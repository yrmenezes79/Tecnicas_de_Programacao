from docx import Document

doc = Document()

# Adicionar um título
doc.add_heading('Título do Documento', level=1)

# Adicionar um parágrafo
doc.add_paragraph('Este é o primeiro parágrafo do documento.')

# Adicionar um parágrafo com estilo diferente
doc.add_paragraph('Este parágrafo está em negrito e itálico.', style='Normal')

# Adicionar uma lista
doc.add_paragraph('Item 1', style='ListBullet')

nome_arquivo = 'exemplo.docx'
doc.save(nome_arquivo)

print(f'Arquivo {nome_arquivo} criado com sucesso!')
